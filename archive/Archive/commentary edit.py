from docx import Document
import os

def load_commentary(doc_path):
    """Extracts all lines that begin with 'Commentary:'"""
    doc = Document(doc_path)
    commentary_lines = [para.text for para in doc.paragraphs if para.text.strip().startswith("Commentary:")]
    return commentary_lines

def edit_commentary(commentary_lines):
    """Prompt user to edit each commentary item"""
    updated_lines = []
    for idx, line in enumerate(commentary_lines, start=1):
        print(f"\n[{idx}] Current: {line}")
        choice = input("Do you want to edit this? (Y/N): ").strip().lower()
        if choice == "y":
            new_line = input("Enter your revised commentary:\n").strip()
            if not new_line.lower().startswith("commentary:"):
                new_line = "Commentary: " + new_line
            updated_lines.append(new_line)
        else:
            updated_lines.append(line)
    return updated_lines

def replace_commentary_in_doc(doc_path, updated_lines):
    """Overwrites commentary lines in the document with updated ones"""
    doc = Document(doc_path)
    new_doc = Document()

    comment_index = 0
    for para in doc.paragraphs:
        if para.text.strip().startswith("Commentary:"):
            if comment_index < len(updated_lines):
                new_doc.add_paragraph(updated_lines[comment_index])
                comment_index += 1
        else:
            new_doc.add_paragraph(para.text)

    new_doc.save(doc_path)
    print(f"\n✅ Commentary updated and saved to {doc_path}")

def main():
    doc_path = "finley_memory.docx"
    if not os.path.exists(doc_path):
        print("❌ File not found:", doc_path)
        return

    print("\n📄 Loading commentary from:", doc_path)
    commentary = load_commentary(doc_path)

    if not commentary:
        print("No commentary found.")
        return

    print(f"\n📝 Found {len(commentary)} commentary items.")
    updated = edit_commentary(commentary)
    replace_commentary_in_doc(doc_path, updated)

if __name__ == "__main__":
    main()